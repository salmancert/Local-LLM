"""Native, always-local tools for common productivity tasks: spreadsheets,
Word documents, PDFs, plain text files, file management, archives (zip/
tar/rar), a calculator, the current date/time, web fetching, and (opt-in)
email and shell commands. These run in-process -- no MCP server needed --
and are merged with any MCP tools (see utils/mcp_manager.py) into one
`tools` list for Foundry Local.

Deeper OS integrations that need a real running application or an OAuth
app registration (Outlook/Excel/Word desktop automation, full browser
automation) aren't implemented here -- they're a better fit for a
dedicated MCP server, which this app already knows how to use. See the
Tools section of README.md for recommended servers.

Safety:
  - File tools are sandboxed to WORKSPACE_DIR (workspace/ in the repo
    root). Every path is resolved and checked to stay inside it before
    any read/write, so a tool call can't read or overwrite files
    elsewhere on the machine, including via `..` traversal.
  - Sending email requires SMTP_HOST/SMTP_USER/SMTP_PASSWORD to be
    configured; the tool isn't even offered to the model otherwise.
  - Shell command execution is off by default. It only becomes available
    if ENABLE_SHELL_TOOL is explicitly set truthy -- see README.md for
    why that's a real decision to make, not a default to accept.

Tool calls are LLM-decided, and the LLM's context can include content
from outside the conversation (fetched web pages, MCP tool output,
uploaded documents) -- treat all of this the same as any other
prompt-injection surface: these tools do what they're told, so the
sandboxing above is the actual safety boundary, not a suggestion.
"""
import ast
import csv
import operator
import os
import re

WORKSPACE_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "workspace")
os.makedirs(WORKSPACE_DIR, exist_ok=True)

try:
    import rarfile as _rarfile_probe  # noqa: F401 -- just checking availability
    _RARFILE_AVAILABLE = True
except ImportError:
    _RARFILE_AVAILABLE = False

MAX_READ_CHARS = 20000
MAX_FETCH_CHARS = 8000

ENABLE_SHELL_TOOL = os.environ.get("ENABLE_SHELL_TOOL", "").strip().lower() in ("1", "true", "yes")
SHELL_TOOL_TIMEOUT = int(os.environ.get("SHELL_TOOL_TIMEOUT", "30"))

SMTP_HOST = os.environ.get("SMTP_HOST")
SMTP_PORT = int(os.environ.get("SMTP_PORT", "587"))
SMTP_USER = os.environ.get("SMTP_USER")
SMTP_PASSWORD = os.environ.get("SMTP_PASSWORD")
_EMAIL_CONFIGURED = bool(SMTP_HOST and SMTP_USER and SMTP_PASSWORD)


def _resolve_path(relative_path):
    """Resolve `relative_path` inside WORKSPACE_DIR, refusing anything
    that would escape it (absolute paths, `..` traversal, symlinked
    escapes). This is the actual sandbox -- every file tool goes through it."""
    if not relative_path or os.path.isabs(relative_path):
        raise ValueError("path must be relative (no leading '/'), e.g. 'notes.txt'")
    full = os.path.realpath(os.path.join(WORKSPACE_DIR, relative_path))
    workspace_real = os.path.realpath(WORKSPACE_DIR)
    if full != workspace_real and not full.startswith(workspace_real + os.sep):
        raise ValueError(f"'{relative_path}' resolves outside the workspace directory")
    return full


def _truncate(text, limit=MAX_READ_CHARS):
    if len(text) > limit:
        return text[:limit] + f"\n[... truncated, {len(text) - limit} more characters]"
    return text


# --- Plain text files (notepad-like) ---------------------------------

def list_workspace_files(subdirectory=""):
    base = _resolve_path(subdirectory) if subdirectory else WORKSPACE_DIR
    if not os.path.isdir(base):
        return f"'{subdirectory}' is not a directory"
    entries = []
    for root, _dirs, files in os.walk(base):
        for name in files:
            rel = os.path.relpath(os.path.join(root, name), WORKSPACE_DIR)
            entries.append(rel)
    return "\n".join(sorted(entries)) if entries else "(workspace is empty)"


def read_text_file(path):
    full = _resolve_path(path)
    with open(full, "r", encoding="utf-8", errors="replace") as f:
        return _truncate(f.read())


def write_text_file(path, content):
    full = _resolve_path(path)
    os.makedirs(os.path.dirname(full), exist_ok=True)
    with open(full, "w", encoding="utf-8") as f:
        f.write(content)
    return f"Wrote {len(content)} characters to {path}"


# --- File management -----------------------------------------------------

def copy_file(source, destination):
    import shutil
    src = _resolve_path(source)
    dst = _resolve_path(destination)
    os.makedirs(os.path.dirname(dst), exist_ok=True)
    shutil.copy2(src, dst)
    return f"Copied {source} to {destination}"


def move_file(source, destination):
    import shutil
    src = _resolve_path(source)
    dst = _resolve_path(destination)
    os.makedirs(os.path.dirname(dst), exist_ok=True)
    shutil.move(src, dst)
    return f"Moved {source} to {destination}"


def delete_file(path):
    full = _resolve_path(path)
    if os.path.isdir(full):
        raise ValueError("path is a directory -- delete_file only removes files")
    os.remove(full)
    return f"Deleted {path}"


# --- PDF ---------------------------------------------------------------

def read_pdf(path):
    from utils.doc_parser import parse_document
    full = _resolve_path(path)
    return _truncate(parse_document(full))


def create_pdf(path, text):
    import fitz  # PyMuPDF
    full = _resolve_path(path)
    os.makedirs(os.path.dirname(full), exist_ok=True)

    doc = fitz.open()
    rect = fitz.paper_rect("letter")
    margin = 50
    text_rect = fitz.Rect(margin, margin, rect.width - margin, rect.height - margin)

    # insert_textbox() returns unused/deficit area as a float, not leftover
    # text, so there's no return value to reflow from -- pre-chunk by a
    # conservative character count per page instead (a full US Letter page
    # at 11pt holds roughly 3000-4000 characters; 2500 leaves headroom).
    chunk_size = 2500
    chunks = [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)] or [""]

    for chunk in chunks:
        page = doc.new_page(width=rect.width, height=rect.height)
        page.insert_textbox(text_rect, chunk, fontsize=11)

    page_count = doc.page_count
    doc.save(full)
    doc.close()
    return f"Wrote PDF with {page_count} page(s) to {path}"


# --- Spreadsheets (CSV / XLSX) -----------------------------------------

def read_spreadsheet(path, sheet_name=None):
    full = _resolve_path(path)
    ext = os.path.splitext(full)[1].lower()

    if ext == ".csv":
        with open(full, "r", encoding="utf-8", errors="replace", newline="") as f:
            rows = list(csv.reader(f))
    elif ext in (".xlsx", ".xlsm"):
        import openpyxl
        wb = openpyxl.load_workbook(full, data_only=True)
        ws = wb[sheet_name] if sheet_name else wb.active
        rows = [[("" if c is None else c) for c in row] for row in ws.iter_rows(values_only=True)]
    else:
        raise ValueError("path must end in .csv, .xlsx, or .xlsm")

    lines = ["\t".join(str(cell) for cell in row) for row in rows]
    return _truncate("\n".join(lines))


def write_spreadsheet(path, rows, sheet_name=None):
    """rows: a list of rows, each a list of cell values."""
    full = _resolve_path(path)
    os.makedirs(os.path.dirname(full), exist_ok=True)
    ext = os.path.splitext(full)[1].lower()

    if ext == ".csv":
        with open(full, "w", encoding="utf-8", newline="") as f:
            csv.writer(f).writerows(rows)
    elif ext in (".xlsx", ".xlsm"):
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        if sheet_name:
            ws.title = _sanitize_sheet_name(sheet_name)
        for row in rows:
            ws.append(row)
        wb.save(full)
    else:
        raise ValueError("path must end in .csv, .xlsx, or .xlsm")

    return f"Wrote {len(rows)} row(s) to {path}"


_INVALID_SHEET_CHARS = re.compile(r"[\[\]:*?/\\]")


def _sanitize_sheet_name(name):
    name = _INVALID_SHEET_CHARS.sub("_", str(name)).strip() or "Sheet"
    return name[:31]  # Excel's sheet-name length limit


def list_spreadsheet_sheets(path):
    import openpyxl
    full = _resolve_path(path)
    wb = openpyxl.load_workbook(full, read_only=True)
    return "\n".join(wb.sheetnames)


def write_spreadsheet_sheet(path, sheet_name, rows):
    """Add or replace one sheet in an .xlsx/.xlsm workbook, preserving its
    other sheets -- unlike write_spreadsheet, which always creates a
    brand-new single-sheet file (so calling it repeatedly on the same
    path just overwrites the previous sheet instead of accumulating
    them). Creates the workbook if it doesn't exist yet."""
    import openpyxl
    full = _resolve_path(path)
    ext = os.path.splitext(full)[1].lower()
    if ext not in (".xlsx", ".xlsm"):
        raise ValueError("write_spreadsheet_sheet only supports .xlsx/.xlsm -- a .csv file has no concept of multiple sheets")

    os.makedirs(os.path.dirname(full), exist_ok=True)
    if os.path.exists(full):
        wb = openpyxl.load_workbook(full)
    else:
        wb = openpyxl.Workbook()
        wb.remove(wb.active)  # drop the default blank sheet; we add our own named one below

    sheet_name = _sanitize_sheet_name(sheet_name)
    if sheet_name in wb.sheetnames:
        del wb[sheet_name]
    ws = wb.create_sheet(title=sheet_name)
    for row in rows:
        ws.append(row)
    wb.save(full)
    return f"Wrote sheet '{sheet_name}' ({len(rows)} row(s)) to {path}"


def split_spreadsheet_by_column(source_path, name_column, output_path=None, has_header=True):
    """Read source_path (.csv/.xlsx/.xlsm), group its rows by the value in
    `name_column` (a header name if has_header, else a 0-based column
    index), and write one sheet per unique value into a single .xlsx
    workbook. This is a deterministic, single-call alternative to asking
    the model to loop write_spreadsheet_sheet once per group -- reliable
    regardless of how well the model handles multi-step tool use."""
    import openpyxl

    full = _resolve_path(source_path)
    ext = os.path.splitext(full)[1].lower()

    if ext == ".csv":
        with open(full, "r", encoding="utf-8", errors="replace", newline="") as f:
            all_rows = list(csv.reader(f))
    elif ext in (".xlsx", ".xlsm"):
        wb_in = openpyxl.load_workbook(full, data_only=True)
        all_rows = [list(row) for row in wb_in.active.iter_rows(values_only=True)]
    else:
        raise ValueError("source_path must end in .csv, .xlsx, or .xlsm")

    if not all_rows:
        raise ValueError("source spreadsheet is empty")

    header = all_rows[0] if has_header else None
    data_rows = all_rows[1:] if has_header else all_rows

    if isinstance(name_column, str):
        if not header:
            raise ValueError("name_column was given as a header name, but has_header is false")
        try:
            col_index = header.index(name_column)
        except ValueError:
            raise ValueError(f"column '{name_column}' not found in header {header}")
    else:
        col_index = int(name_column)

    groups = {}
    for row in data_rows:
        key = str(row[col_index]) if col_index < len(row) and row[col_index] is not None else ""
        groups.setdefault(key, []).append(row)

    if not output_path:
        base, _ = os.path.splitext(source_path)
        output_path = f"{base}_split.xlsx"
    out_full = _resolve_path(output_path)
    if not out_full.endswith((".xlsx", ".xlsm")):
        raise ValueError("output_path must end in .xlsx or .xlsm -- only these formats hold multiple sheets in one file")

    os.makedirs(os.path.dirname(out_full), exist_ok=True)
    wb_out = openpyxl.Workbook()
    wb_out.remove(wb_out.active)
    used_names = set()
    sheet_names = []
    for name, rows in groups.items():
        sheet_name = _sanitize_sheet_name(name)
        # Two different raw values can sanitize to the same name (e.g.
        # "A:B" and "A/B" both become "A_B") -- dedupe explicitly rather
        # than relying on openpyxl's own auto-rename, so the name we
        # report back always matches the sheet actually created.
        base_name, suffix = sheet_name, 2
        while sheet_name in used_names:
            sheet_name = f"{base_name[:29]}_{suffix}"
            suffix += 1
        used_names.add(sheet_name)
        sheet_names.append(sheet_name)

        ws = wb_out.create_sheet(title=sheet_name)
        if header:
            ws.append(header)
        for row in rows:
            ws.append(row)
    wb_out.save(out_full)

    return f"Created {len(groups)} sheet(s) in {output_path}: {', '.join(sorted(sheet_names))}"


# --- Word documents (.docx) --------------------------------------------

def read_word_document(path):
    import docx
    full = _resolve_path(path)
    document = docx.Document(full)
    return _truncate("\n".join(p.text for p in document.paragraphs))


def write_word_document(path, content):
    """content: the document text; blank lines start a new paragraph."""
    import docx
    full = _resolve_path(path)
    os.makedirs(os.path.dirname(full), exist_ok=True)

    document = docx.Document()
    for paragraph in content.split("\n"):
        document.add_paragraph(paragraph)
    document.save(full)
    return f"Wrote Word document to {path}"


# --- Web browsing (fetch + extract readable text) ----------------------

def fetch_webpage(url):
    import requests
    from bs4 import BeautifulSoup

    if not re.match(r"^https?://", url, re.IGNORECASE):
        raise ValueError("url must start with http:// or https://")

    response = requests.get(url, timeout=15, headers={"User-Agent": "Iris-local-assistant/1.0"})
    response.raise_for_status()

    soup = BeautifulSoup(response.text, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "noscript"]):
        tag.decompose()
    text = " ".join(soup.get_text(separator=" ").split())
    return _truncate(text, MAX_FETCH_CHARS)


# --- Archives (zip / tar / rar) -----------------------------------------

def _iter_files_for_archive(paths):
    """Yield (real_filesystem_path, archive_member_name) pairs for a list
    of workspace-relative paths, expanding any directories recursively."""
    for p in paths:
        resolved = _resolve_path(p)
        if os.path.isdir(resolved):
            for root, _dirs, filenames in os.walk(resolved):
                for name in filenames:
                    full_file = os.path.join(root, name)
                    arcname = os.path.relpath(full_file, WORKSPACE_DIR)
                    yield full_file, arcname
        else:
            yield resolved, p


def _check_no_archive_escape(names, dest_dir):
    """Guard against "zip-slip": an archive member whose name (e.g.
    '../../etc/cron.d/x' or an absolute path) would land outside the
    intended destination once extracted. Both zipfile and tarfile will
    happily write wherever a malicious member name points unless the
    caller checks first -- so we check first."""
    dest_real = os.path.realpath(dest_dir)
    for name in names:
        target = os.path.realpath(os.path.join(dest_dir, name))
        if target != dest_real and not target.startswith(dest_real + os.sep):
            raise ValueError(f"archive member '{name}' would extract outside the destination -- refusing to extract")


def create_zip_archive(path, files):
    import zipfile
    full = _resolve_path(path)
    os.makedirs(os.path.dirname(full), exist_ok=True)
    count = 0
    with zipfile.ZipFile(full, "w", zipfile.ZIP_DEFLATED) as zf:
        for full_file, arcname in _iter_files_for_archive(files):
            zf.write(full_file, arcname=arcname)
            count += 1
    return f"Created {path} with {count} file(s)"


def extract_zip_archive(path, destination=""):
    import zipfile
    full = _resolve_path(path)
    dest_dir = _resolve_path(destination) if destination else WORKSPACE_DIR
    os.makedirs(dest_dir, exist_ok=True)
    with zipfile.ZipFile(full, "r") as zf:
        names = zf.namelist()
        _check_no_archive_escape(names, dest_dir)
        zf.extractall(dest_dir)
    return f"Extracted {len(names)} file(s) from {path} to {destination or '.'}"


def create_tar_archive(path, files):
    import tarfile
    full = _resolve_path(path)
    os.makedirs(os.path.dirname(full), exist_ok=True)
    if full.endswith((".tar.gz", ".tgz")):
        mode = "w:gz"
    elif full.endswith((".tar.bz2", ".tbz2")):
        mode = "w:bz2"
    else:
        mode = "w"
    count = 0
    with tarfile.open(full, mode) as tf:
        for full_file, arcname in _iter_files_for_archive(files):
            tf.add(full_file, arcname=arcname)
            count += 1
    return f"Created {path} with {count} file(s)"


def extract_tar_archive(path, destination=""):
    import tarfile
    full = _resolve_path(path)
    dest_dir = _resolve_path(destination) if destination else WORKSPACE_DIR
    os.makedirs(dest_dir, exist_ok=True)
    with tarfile.open(full, "r:*") as tf:
        members = tf.getmembers()
        _check_no_archive_escape([m.name for m in members], dest_dir)
        tf.extractall(dest_dir, members=members)
    return f"Extracted {len(members)} file(s) from {path} to {destination or '.'}"


def extract_rar_archive(path, destination=""):
    # RAR is a proprietary format with no free encoder, so only extraction
    # is offered (no create_rar_archive). Needs the `rarfile` package AND
    # a system unrar/unar/bsdtar binary on PATH to actually decompress --
    # rarfile is just a wrapper around one of those.
    import rarfile
    full = _resolve_path(path)
    dest_dir = _resolve_path(destination) if destination else WORKSPACE_DIR
    os.makedirs(dest_dir, exist_ok=True)
    with rarfile.RarFile(full) as rf:
        names = rf.namelist()
        _check_no_archive_escape(names, dest_dir)
        rf.extractall(dest_dir)
    return f"Extracted {len(names)} file(s) from {path} to {destination or '.'}"


# --- Utility: calculator, current date/time ------------------------------

_SAFE_OPERATORS = {
    ast.Add: operator.add,
    ast.Sub: operator.sub,
    ast.Mult: operator.mul,
    ast.Div: operator.truediv,
    ast.FloorDiv: operator.floordiv,
    ast.Mod: operator.mod,
    ast.Pow: operator.pow,
    ast.USub: operator.neg,
    ast.UAdd: operator.pos,
}


def _eval_arithmetic_node(node):
    """Evaluate one node of an arithmetic-only AST. Deliberately not
    `eval()`: this whitelist accepts numeric literals and +-*/%** only --
    no names, no calls, no attribute/subscript access -- so it can't be
    used to run arbitrary code no matter what expression an LLM passes in."""
    if isinstance(node, ast.Constant) and isinstance(node.value, (int, float)):
        return node.value
    if isinstance(node, ast.BinOp) and type(node.op) in _SAFE_OPERATORS:
        return _SAFE_OPERATORS[type(node.op)](_eval_arithmetic_node(node.left), _eval_arithmetic_node(node.right))
    if isinstance(node, ast.UnaryOp) and type(node.op) in _SAFE_OPERATORS:
        return _SAFE_OPERATORS[type(node.op)](_eval_arithmetic_node(node.operand))
    raise ValueError("only numbers and + - * / // % ** are allowed")


def calculate(expression):
    tree = ast.parse(expression, mode="eval")
    return str(_eval_arithmetic_node(tree.body))


def get_current_datetime():
    import datetime
    return datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S %A")


# --- Email (opt-in: only offered if SMTP_* is configured) --------------

def send_email(to, subject, body):
    if not _EMAIL_CONFIGURED:
        raise RuntimeError("Email isn't configured -- set SMTP_HOST, SMTP_USER, and SMTP_PASSWORD.")

    import smtplib
    from email.mime.text import MIMEText

    msg = MIMEText(body)
    msg["Subject"] = subject
    msg["From"] = SMTP_USER
    msg["To"] = to

    with smtplib.SMTP(SMTP_HOST, SMTP_PORT, timeout=15) as server:
        server.starttls()
        server.login(SMTP_USER, SMTP_PASSWORD)
        server.sendmail(SMTP_USER, [to], msg.as_string())
    return f"Email sent to {to}"


# --- Shell / PowerShell (opt-in: off unless ENABLE_SHELL_TOOL is set) --

def run_shell_command(command):
    if not ENABLE_SHELL_TOOL:
        raise RuntimeError("Shell command execution is disabled -- set ENABLE_SHELL_TOOL=true to enable it (see README.md for what that means).")

    import subprocess
    if os.name == "nt":
        argv = ["powershell", "-NoProfile", "-NonInteractive", "-Command", command]
    else:
        argv = ["/bin/sh", "-c", command]

    result = subprocess.run(
        argv, cwd=WORKSPACE_DIR, capture_output=True, text=True, timeout=SHELL_TOOL_TIMEOUT,
    )
    output = (result.stdout or "") + (result.stderr or "")
    return _truncate(f"[exit code {result.returncode}]\n{output}", 4000)


# --- Tool registration ---------------------------------------------------

_STATIC_TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "local__list_workspace_files",
            "description": "List files in the local workspace (the sandboxed folder these tools read/write in).",
            "parameters": {
                "type": "object",
                "properties": {"subdirectory": {"type": "string", "description": "Optional subdirectory to list, relative to the workspace root"}},
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "local__read_text_file",
            "description": "Read a plain text file (like Notepad) from the local workspace.",
            "parameters": {
                "type": "object",
                "properties": {"path": {"type": "string", "description": "File path relative to the workspace root, e.g. 'notes.txt'"}},
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "local__write_text_file",
            "description": "Write (create or overwrite) a plain text file in the local workspace.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "File path relative to the workspace root, e.g. 'notes.txt'"},
                    "content": {"type": "string", "description": "The full text content to write"},
                },
                "required": ["path", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "local__read_pdf",
            "description": "Extract the text content of a PDF file in the local workspace.",
            "parameters": {
                "type": "object",
                "properties": {"path": {"type": "string", "description": "PDF path relative to the workspace root"}},
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "local__create_pdf",
            "description": "Create a new PDF file containing the given text, in the local workspace.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "PDF path to create, relative to the workspace root"},
                    "text": {"type": "string", "description": "The text content of the PDF"},
                },
                "required": ["path", "text"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "local__read_spreadsheet",
            "description": "Read a spreadsheet (.csv, .xlsx, or .xlsm) from the local workspace and return its rows as tab-separated text.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Spreadsheet path relative to the workspace root"},
                    "sheet_name": {"type": "string", "description": "Sheet name for .xlsx files (optional; defaults to the active sheet)"},
                },
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "local__write_spreadsheet",
            "description": "Create a new spreadsheet (.csv, .xlsx, or .xlsm) with one sheet of data, in the local workspace. This always creates a fresh single-sheet file -- calling it again on the same path replaces the whole file. To build a workbook with multiple named sheets, use local__write_spreadsheet_sheet instead.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Spreadsheet path to create, relative to the workspace root"},
                    "rows": {
                        "type": "array",
                        "description": "Rows of data; each row is an array of cell values",
                        "items": {"type": "array", "items": {}},
                    },
                    "sheet_name": {"type": "string", "description": "Sheet name for .xlsx files (optional)"},
                },
                "required": ["path", "rows"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "local__write_spreadsheet_sheet",
            "description": "Add or replace one sheet in an .xlsx/.xlsm workbook, keeping its other sheets intact. Creates the workbook if it doesn't exist. Use this to build up a workbook with multiple named sheets, e.g. calling it once per name.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Workbook path (.xlsx/.xlsm), relative to the workspace root"},
                    "sheet_name": {"type": "string", "description": "Name of the sheet to add or replace"},
                    "rows": {
                        "type": "array",
                        "description": "Rows of data; each row is an array of cell values",
                        "items": {"type": "array", "items": {}},
                    },
                },
                "required": ["path", "sheet_name", "rows"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "local__list_spreadsheet_sheets",
            "description": "List the sheet names in an .xlsx/.xlsm workbook.",
            "parameters": {
                "type": "object",
                "properties": {"path": {"type": "string", "description": "Workbook path relative to the workspace root"}},
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "local__split_spreadsheet_by_column",
            "description": "Split a spreadsheet into multiple sheets in one new workbook, one sheet per unique value in a chosen column (e.g. one sheet per person or category, each containing that group's rows). Does the whole split in one call -- more reliable than looping local__write_spreadsheet_sheet per value.",
            "parameters": {
                "type": "object",
                "properties": {
                    "source_path": {"type": "string", "description": "Source spreadsheet (.csv, .xlsx, .xlsm), relative to the workspace root"},
                    "name_column": {"description": "Column to split by: the header name if has_header is true, otherwise a 0-based column index"},
                    "output_path": {"type": "string", "description": "Output .xlsx path, relative to the workspace root (optional; defaults to '<source>_split.xlsx')"},
                    "has_header": {"type": "boolean", "description": "Whether the first row is a header row (default true)"},
                },
                "required": ["source_path", "name_column"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "local__read_word_document",
            "description": "Read the text content of a Word document (.docx) in the local workspace.",
            "parameters": {
                "type": "object",
                "properties": {"path": {"type": "string", "description": "Document path relative to the workspace root"}},
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "local__write_word_document",
            "description": "Create a new Word document (.docx) with the given text in the local workspace. Blank lines start a new paragraph.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Document path to create, relative to the workspace root"},
                    "content": {"type": "string", "description": "The text content of the document"},
                },
                "required": ["path", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "local__fetch_webpage",
            "description": "Fetch a web page and return its readable text content (for web browsing / research).",
            "parameters": {
                "type": "object",
                "properties": {"url": {"type": "string", "description": "The URL to fetch, including http:// or https://"}},
                "required": ["url"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "local__copy_file",
            "description": "Copy a file within the local workspace.",
            "parameters": {
                "type": "object",
                "properties": {
                    "source": {"type": "string", "description": "Source path relative to the workspace root"},
                    "destination": {"type": "string", "description": "Destination path relative to the workspace root"},
                },
                "required": ["source", "destination"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "local__move_file",
            "description": "Move or rename a file within the local workspace.",
            "parameters": {
                "type": "object",
                "properties": {
                    "source": {"type": "string", "description": "Source path relative to the workspace root"},
                    "destination": {"type": "string", "description": "Destination path relative to the workspace root"},
                },
                "required": ["source", "destination"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "local__delete_file",
            "description": "Delete a file (not a directory) in the local workspace.",
            "parameters": {
                "type": "object",
                "properties": {"path": {"type": "string", "description": "File path relative to the workspace root"}},
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "local__create_zip_archive",
            "description": "Create a .zip archive in the local workspace from one or more files/folders.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Zip file path to create, relative to the workspace root"},
                    "files": {"type": "array", "items": {"type": "string"}, "description": "Files or folders to include, relative to the workspace root"},
                },
                "required": ["path", "files"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "local__extract_zip_archive",
            "description": "Extract a .zip archive in the local workspace.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Zip file path relative to the workspace root"},
                    "destination": {"type": "string", "description": "Folder to extract into, relative to the workspace root (optional; defaults to the workspace root)"},
                },
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "local__create_tar_archive",
            "description": "Create a .tar, .tar.gz/.tgz, or .tar.bz2 archive in the local workspace from one or more files/folders.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Archive path to create, relative to the workspace root -- extension picks the compression (.tar, .tar.gz, .tgz, .tar.bz2)"},
                    "files": {"type": "array", "items": {"type": "string"}, "description": "Files or folders to include, relative to the workspace root"},
                },
                "required": ["path", "files"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "local__extract_tar_archive",
            "description": "Extract a .tar, .tar.gz/.tgz, or .tar.bz2 archive in the local workspace.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Archive path relative to the workspace root"},
                    "destination": {"type": "string", "description": "Folder to extract into, relative to the workspace root (optional; defaults to the workspace root)"},
                },
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "local__calculate",
            "description": "Evaluate an arithmetic expression (+ - * / // % **) and return the result.",
            "parameters": {
                "type": "object",
                "properties": {"expression": {"type": "string", "description": "e.g. '(12 + 8) * 3'"}},
                "required": ["expression"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "local__get_current_datetime",
            "description": "Get the current local date and time.",
            "parameters": {"type": "object", "properties": {}},
        },
    },
]

_RAR_TOOL = {
    "type": "function",
    "function": {
        "name": "local__extract_rar_archive",
        "description": "Extract a .rar archive in the local workspace. (Extraction only -- RAR is a proprietary format with no free encoder, so there's no way to create one.)",
        "parameters": {
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "RAR file path relative to the workspace root"},
                "destination": {"type": "string", "description": "Folder to extract into, relative to the workspace root (optional; defaults to the workspace root)"},
            },
            "required": ["path"],
        },
    },
}

_EMAIL_TOOL = {
    "type": "function",
    "function": {
        "name": "local__send_email",
        "description": "Send an email via the configured SMTP account.",
        "parameters": {
            "type": "object",
            "properties": {
                "to": {"type": "string", "description": "Recipient email address"},
                "subject": {"type": "string", "description": "Email subject"},
                "body": {"type": "string", "description": "Email body text"},
            },
            "required": ["to", "subject", "body"],
        },
    },
}

_SHELL_TOOL = {
    "type": "function",
    "function": {
        "name": "local__run_shell_command",
        "description": "Run a shell command (PowerShell on Windows, sh elsewhere) in the local workspace directory and return its output.",
        "parameters": {
            "type": "object",
            "properties": {"command": {"type": "string", "description": "The command to run"}},
            "required": ["command"],
        },
    },
}

_HANDLERS = {
    "local__list_workspace_files": lambda a: list_workspace_files(a.get("subdirectory", "")),
    "local__read_text_file": lambda a: read_text_file(a["path"]),
    "local__write_text_file": lambda a: write_text_file(a["path"], a["content"]),
    "local__copy_file": lambda a: copy_file(a["source"], a["destination"]),
    "local__move_file": lambda a: move_file(a["source"], a["destination"]),
    "local__delete_file": lambda a: delete_file(a["path"]),
    "local__read_pdf": lambda a: read_pdf(a["path"]),
    "local__create_pdf": lambda a: create_pdf(a["path"], a["text"]),
    "local__read_spreadsheet": lambda a: read_spreadsheet(a["path"], a.get("sheet_name")),
    "local__write_spreadsheet": lambda a: write_spreadsheet(a["path"], a["rows"], a.get("sheet_name")),
    "local__write_spreadsheet_sheet": lambda a: write_spreadsheet_sheet(a["path"], a["sheet_name"], a["rows"]),
    "local__list_spreadsheet_sheets": lambda a: list_spreadsheet_sheets(a["path"]),
    "local__split_spreadsheet_by_column": lambda a: split_spreadsheet_by_column(
        a["source_path"], a["name_column"], a.get("output_path"), a.get("has_header", True)
    ),
    "local__read_word_document": lambda a: read_word_document(a["path"]),
    "local__write_word_document": lambda a: write_word_document(a["path"], a["content"]),
    "local__fetch_webpage": lambda a: fetch_webpage(a["url"]),
    "local__create_zip_archive": lambda a: create_zip_archive(a["path"], a["files"]),
    "local__extract_zip_archive": lambda a: extract_zip_archive(a["path"], a.get("destination", "")),
    "local__create_tar_archive": lambda a: create_tar_archive(a["path"], a["files"]),
    "local__extract_tar_archive": lambda a: extract_tar_archive(a["path"], a.get("destination", "")),
    "local__extract_rar_archive": lambda a: extract_rar_archive(a["path"], a.get("destination", "")),
    "local__calculate": lambda a: calculate(a["expression"]),
    "local__get_current_datetime": lambda a: get_current_datetime(),
    "local__send_email": lambda a: send_email(a["to"], a["subject"], a["body"]),
    "local__run_shell_command": lambda a: run_shell_command(a["command"]),
}


def get_tool_schemas():
    tools = list(_STATIC_TOOLS)
    if _RARFILE_AVAILABLE:
        tools.append(_RAR_TOOL)
    if _EMAIL_CONFIGURED:
        tools.append(_EMAIL_TOOL)
    if ENABLE_SHELL_TOOL:
        tools.append(_SHELL_TOOL)
    return tools


def is_local_tool(qualified_name):
    return qualified_name in _HANDLERS


def call_tool(qualified_name, arguments):
    handler = _HANDLERS.get(qualified_name)
    if handler is None:
        return f"Error: unknown tool '{qualified_name}'"
    try:
        return str(handler(arguments))
    except Exception as e:
        return f"Error: {e}"
